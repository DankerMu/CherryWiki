from __future__ import annotations

from pathlib import Path

from docx import Document

from .base_parser import BaseParser, ParseResult
from .utils import compact_blank_lines, markdown_table, package_version


class DocxParser(BaseParser):
    source_type = "docx"
    extraction_tool = "python-docx"
    extraction_version = package_version("python-docx")

    def parse(self, file_path: Path) -> ParseResult:
        document = Document(file_path)
        sections: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            sections.append(self._format_paragraph(text, paragraph.style.name if paragraph.style else ""))

        for table in document.tables:
            table_markdown = markdown_table([[cell.text.strip() for cell in row.cells] for row in table.rows])
            if table_markdown:
                sections.append(table_markdown)

        content = compact_blank_lines("\n\n".join(sections))
        return ParseResult(
            content=content,
            metadata={
                "source_type": self.source_type,
                "extraction_tool": self.extraction_tool,
                "extraction_version": self.extraction_version,
                "extraction_params": {"preserve_headings": True, "extract_tables": True},
                "page_count": 1,
                "char_count": len(content),
                "image_count": len(document.inline_shapes),
            },
        )

    def _format_paragraph(self, text: str, style_name: str) -> str:
        if style_name.startswith("Heading "):
            try:
                level = max(1, min(6, int(style_name.removeprefix("Heading ").strip())))
            except ValueError:
                level = 2
            return f"{'#' * level} {text}"
        if "List Bullet" in style_name:
            return f"- {text}"
        if "List Number" in style_name:
            return f"1. {text}"
        return text
